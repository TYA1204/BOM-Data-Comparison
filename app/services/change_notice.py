"""
整机清机更改通知单 — 自动填表工具
直接操作模版表格，在「更改内容」行后插入差异明细行
"""

import os
import re
import sqlite3
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK


# ── Constants ──────────────────────────────────────────────
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

TEMPLATE_PATH = os.path.join(PROJECT_ROOT, '整机清机更改通知单.docx')
OUTPUT_DIR = os.path.join(PROJECT_ROOT, 'reports')

# Font fallback list: Windows → Linux → macOS
# 宋体/SimSun (Windows) | Noto Sans CJK SC (Linux) | Arial Unicode MS (macOS)
FONT_FALLBACK = ['宋体', 'SimSun', 'Noto Sans CJK SC', 'WenQuanYi Micro Hei', 'Microsoft YaHei', 'Arial Unicode MS']

DIFF_LABELS = {'added': '新增', 'removed': '删除', 'modified': '变更'}

# ── Colors ──────────────────────────────────────────────────
GREEN      = '16A34A'
RED        = 'DC2626'
AMBER      = 'CA8A04'
GREEN_BG   = 'DCFCE7'
RED_BG     = 'FEF2F2'
AMBER_BG   = 'FEF9C3'
HEADER_BG  = '1E40AF'
HEADER_FG  = 'FFFFFF'
BLACK      = '333333'
GRAY       = '999999'


# ── Data Helpers ───────────────────────────────────────────

def _set_font(run, font_list=None):
    """Apply font with fallback list.

    python-docx only sets the font name; actual font substitution
    is handled by Word/Office when the requested font is unavailable.
    We set both the ASCII (name) and East Asian (rFonts element) to
    maximize compatibility across Windows / macOS / Linux.
    """
    if font_list is None:
        font_list = FONT_FALLBACK
    run.font.name = font_list[0]
    # Also set East Asian font explicitly for CJK text
    rPr = run._element.getparent()
    if rPr is not None:
        from lxml import etree
        rFonts = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + attr, font_list[0])


def _g(d, key, default=None):
    try:
        return d[key]
    except (KeyError, IndexError):
        return default


def clean_material_name(name):
    """Extract core material name by stripping specs and noise.

    Strategy: scan tokens until we hit a digit-starting or special-char
    token (spec boundary), keeping only the core identifier. A tiny
    blacklist handles universal noise words (brands, colors, certs).
    """
    import re

    if not name:
        return ""

    # Step 1: strip parenthesized content
    name = re.sub(r"[（【].*[）】]", " ", name)
    name = re.sub(r"\([^)]*\)", " ", name)
    # Step 1b: strip ALT/BOM metadata markers (| #ALT#..., | ECN, | REACH, etc.)
    name = re.sub(r'\s*\|\s*#.*$', '', name)
    name = re.sub(r'\s*\|\s*(ECN|REACH|RoHS|NCN\d+).*$', '', name)
    # Step 1c: simplify verbose component type prefixes
    name = re.sub(r'插件交流安规(薄膜|陶瓷)电容', r'安规\1电容', name)
    name = re.sub(r'贴片低容陶瓷电容', '陶瓷电容', name)
    name = re.sub(r'贴片高容陶瓷电容', '陶瓷电容', name)
    name = re.sub(r'贴片射频陶瓷电容', '射频电容', name)
    name = re.sub(r'高频低阻电解电容', '电解电容', name)
    name = re.sub(r'高压铝电解电容', '电解电容', name)
    name = re.sub(r'贴片铝电解电容', '铝电解电容', name)
    name = re.sub(r'中低压MOS管500V以下', 'MOS管', name)
    # Step 1d: strip packaging suffixes
    name = re.sub(r'\b卧式成型\b', '', name)
    name = re.sub(r'\b立式成型\b', '', name)
    name = re.sub(r'\b直脚编带\b', '', name)
    name = re.sub(r'\b安全膜\b', '', name)
    name = re.sub(r'\b非安全膜\b', '', name)
    tokens = name.split()

    # Universal noise — only words that are never part of a core material name
    SKIP = {
        # Brands / regions
        "SKYWORTH", "coocaa", "COOCAA", "HKC", "WANZHONG", "wanzhong", "WZ",
        "Shenzhen", "CHUANGWEI-RGB", "Electro", "深圳创维RGB电子",
        "内销", "通用", "中国", "无",
        # Certifications
        "RoHS", "REACH", "CBU",
        # Colors
        "黑色", "白色", "灰色", "本色", "中黑", "磨砂黑",
        # Materials / forms
        "PU", "HAIMIAN", "EVA", "PET", "PE", "HIPS", "PC+ABS", "SECC",
        "片料", "规则", "不规则",
        # Process / tech
        "不加硬", "镀黑锌", "镀锌", "镀镍", "镀锡", "镀银", "白锌", "彩锌",
        "加硬", "硬化", "热处理", "环保", "环保型",
        # Screw descriptors
        "十字", "一字", "内六角", "外六角", "梅花",
        "工业牙", "自攻牙", "机牙", "自攻",
        "平尾", "尖尾", "锥尾", "割尾",
        "全螺纹", "半螺纹",
        "无垫圈", "有垫圈", "带垫圈", "含垫圈",
        "无弹垫", "有弹垫", "无平垫", "有平垫",
        "非耐落", "耐落",
        # Directional
        "下", "上", "左", "右",
        "INSIDE", "BOTTOM", "BOX",
        # Filler
        "Pro", "for", "and", "the", "with", "width",
        "防静电", "耐高温",
        "非REACH", "非无卤", "非防静电", "非阻燃", "非RoHS", "无耐温要求", "非卤",
        "彩色印刷", "A.彩色", "2P", "NEW", "EMMC", "64G", "PDM",
        "高频膜", "单面", "邮票机贴", "非高频膜",
        "研发机贴", "研发手插", "研发机插",
        "fpc", "FPC", "fengwoban",
        "附接订单", "此订单。",
    }

    # Electronic component value units — these ARE core identity (resistors, caps, etc.)
    _comp_unit_re = re.compile(
        r'[\u03BCu]?[FfHh]$'        # μF, nF, pF, μH, nH
        r'|[KMkm]?[Ω]$'            # Ω, KΩ, MΩ
        r'|[VvAa]$'                # V, A, mA
    )

    kept = []
    value_count = 0
    for t in tokens:
        if t in SKIP:
            continue
        # Special-char token with no Chinese → spec boundary
        if re.search(r"[*~/±°→×]", t) and not re.search(r"[\u4e00-\u9fff]", t):
            break
        # Starts with digit
        if re.match(r"^\d", t):
            # Model code (100A5H, 8R713) → keep then stop
            if re.match(r"^\d+[A-Za-z]+\d", t):
                kept.append(t)
                break
            # Type code (4T, 10A, 500V-like) — keep, continue for values
            # Only short codes: max 3 digits + 1-2 letters, and NOT ending in common units
            if re.match(r"^\d{1,3}[A-Za-z]{1,2}$", t) and not re.match(r'.*[mMVVA]$', t):
                kept.append(t)
                continue
            # Component value: 10KΩ, 240Ω, 100nF, 4.7μH, 50V, 650mA
            if _comp_unit_re.search(t) and value_count < 2:
                # Skip reverse leakage current specs (multi-million μA)
                if re.search(r'[\u03BCu]A$', t):
                    val_match = re.match(r'([\d.]+)', t)
                    if val_match and float(val_match.group(1)) > 100000:
                        break
                kept.append(t)
                value_count += 1
                continue
            break
        # Pure codes (T20, BOX, USB3.0, fengwoban, S-PCB, Hanger, etc.)
        if re.fullmatch(r"[A-Z0-9_.-]{2,}", t):
            continue
        if re.fullmatch(r"[A-Z][a-z]+", t):
            continue
        if re.fullmatch(r"[a-z]{3,}", t):
            continue
        if re.fullmatch(r"[A-Z]-[A-Z]+", t):
            continue
        kept.append(t)

    # Max 5 tokens fallback
    if len(kept) > 5:
        kept = kept[:5]

    return " ".join(kept).strip()


def _clean_ref_text(ref_text):
    """Remove signature placeholder and other template noise from reference text."""
    if not ref_text:
        return ref_text
    # Strip signature placeholders: 批 准: __________ 审 核: __________ 拟 制: __________
    ref_text = re.sub(r'\s*(?:批\s*准|审\s*核|拟\s*制)[:：]\s*_{3,}.*$', '', ref_text)
    return ref_text.strip()


def get_diff_rows(conn, task_id, source_bom_id=None, target_bom_id=None):
    """Get all diff rows for the task, ready for table insertion.

    When source_bom_id/target_bom_id are provided, parent component names are
    looked up from the correct BOM: source BOM for DEL items, target BOM for
    ADD items.  This prevents cross-BOM name pollution when the same PN has
    different names in different BOMs.
    """
    diffs = conn.execute(
        '''SELECT * FROM comparison_result WHERE task_id=?
           ORDER BY CASE diff_type
               WHEN 'added' THEN 1 WHEN 'removed' THEN 2 ELSE 3
           END, COALESCE(line_no_b, line_no_a)''',
        (task_id,)
    ).fetchall()

    # ── Build BOM-aware parent name lookups ──
    parent_pns_src = set()  # parent_pn_a: source BOM side
    parent_pns_tgt = set()  # parent_pn_b: target BOM side
    for d in diffs:
        pa = (_g(d, 'parent_pn_a') or '').strip()
        pb = (_g(d, 'parent_pn_b') or '').strip()
        if pa: parent_pns_src.add(pa)
        if pb: parent_pns_tgt.add(pb)

    parent_lookup_src = {}
    parent_lookup_tgt = {}

    if source_bom_id and parent_pns_src:
        placeholders = ','.join(['?' for _ in parent_pns_src])
        pns = conn.execute(
            f'''SELECT part_number, part_name FROM bom_item
                WHERE bom_id=? AND part_number IN ({placeholders})''',
            [source_bom_id] + list(parent_pns_src)
        ).fetchall()
        for pn_row in pns:
            raw_name = (pn_row['part_name'] or '\u90e8\u4ef6').strip()
            parent_lookup_src[pn_row['part_number']] = clean_material_name(raw_name)

    if target_bom_id and parent_pns_tgt:
        placeholders = ','.join(['?' for _ in parent_pns_tgt])
        pns = conn.execute(
            f'''SELECT part_number, part_name FROM bom_item
                WHERE bom_id=? AND part_number IN ({placeholders})''',
            [target_bom_id] + list(parent_pns_tgt)
        ).fetchall()
        for pn_row in pns:
            raw_name = (pn_row['part_name'] or '\u90e8\u4ef6').strip()
            parent_lookup_tgt[pn_row['part_number']] = clean_material_name(raw_name)

    # Fallback: if bom_ids not provided, query all BOMs (backward compatible)
    if not source_bom_id and not target_bom_id:
        all_parent_pns = parent_pns_src | parent_pns_tgt
        if all_parent_pns:
            placeholders = ','.join(['?' for _ in all_parent_pns])
            pns = conn.execute(
                f'SELECT part_number, part_name FROM bom_item WHERE part_number IN ({placeholders})',
                list(all_parent_pns)
            ).fetchall()
            for pn_row in pns:
                raw_name = (pn_row['part_name'] or '\u90e8\u4ef6').strip()
                cleaned = clean_material_name(raw_name)
                parent_lookup_src[pn_row['part_number']] = cleaned
                parent_lookup_tgt[pn_row['part_number']] = cleaned

    rows = []
    for d in diffs:
        dt = d['diff_type']
        dc = d['diff_category']

        if dt == 'added':
            pn = (_g(d, 'part_number_b') or '').strip()
            nm = clean_material_name((_g(d, 'part_name_b') or '').strip())
            qty = _g(d, 'quantity_b', 1)
            qty_str = str(int(qty)) if qty is not None else '1'
            parent_pn = (_g(d, 'parent_pn_b') or '').strip()
            # ADD = P3EM side → prefer target BOM name, fallback to source
            parent_name = parent_lookup_tgt.get(parent_pn, parent_lookup_src.get(parent_pn, ''))
            type_label = 'ADD'
            line_no = d['line_no_b'] or d['line_no_a'] or 0
            ref_val = (_g(d, 'reference_b') or _g(d, 'reference_a') or '').strip()
            ref_val = _clean_ref_text(ref_val)
        elif dt == 'removed':
            pn = (_g(d, 'part_number_a') or '').strip()
            nm = clean_material_name((_g(d, 'part_name_a') or '').strip())
            qty = _g(d, 'quantity_a', 1)
            qty_str = str(int(qty)) if qty is not None else '1'
            parent_pn = (_g(d, 'parent_pn_a') or '').strip()
            # DEL = H5F side → prefer source BOM name, fallback to target
            parent_name = parent_lookup_src.get(parent_pn, parent_lookup_tgt.get(parent_pn, ''))
            type_label = 'DEL'
            line_no = d['line_no_a'] or d['line_no_b'] or 0
            ref_val = (_g(d, 'reference_a') or _g(d, 'reference_b') or '').strip()
            ref_val = _clean_ref_text(ref_val)
        else:
            pn = (_g(d, 'part_number_a') or _g(d, 'part_number_b') or '').strip()
            nm = clean_material_name((_g(d, 'part_name_a') or _g(d, 'part_name_b') or '').strip())
            old_qty = str(_g(d, 'old_value', '')).strip()
            new_qty = str(_g(d, 'new_value', '')).strip()
            parent_pn = (_g(d, 'parent_pn_a') or _g(d, 'parent_pn_b') or '').strip()
            # MOD → prefer target, fallback to source
            parent_name = parent_lookup_tgt.get(parent_pn, parent_lookup_src.get(parent_pn, ''))
            type_label = 'MOD'
            line_no = d['line_no_a'] or d['line_no_b'] or 0
            # MOD ref → 优先用目标 BOM 位号（当前状态），源位号做回退
            ref_val = (_g(d, 'reference_b') or _g(d, 'reference_a') or '').strip()
            ref_val = _clean_ref_text(ref_val)

        row_data = {
            'pn': pn, 'name': nm,
            'parent_pn': parent_pn,
            'parent_name': parent_name,
            'type': dt, 'type_label': type_label,
            'line_no': line_no,
            'ref': ref_val,
        }
        if dt == 'modified':
            row_data['old_qty'] = old_qty
            row_data['new_qty'] = new_qty
            row_data['diff_category'] = _g(d, 'diff_category', '')
        else:
            row_data['qty'] = qty_str
        rows.append(row_data)
    return rows


def _get_functional_key(comp_name):
    """Extract functional prefix for cross-model (H5F→P3EM) parent matching.

    Returns the first meaningful word(s) that identify the component's role,
    ignoring the model-specific suffix.  Examples:
      '电子BOM 100H5FP'    → '电子BOM'
      '包装组件 100P3E MAX' → '包装组件'
      '液晶模组 99.5inc...' → '液晶模组'
      'LED二合一电源组件...' → 'LED二合一电源组件'
    """
    if not comp_name:
        return ''
    parts = comp_name.split()
    if not parts:
        return ''
    # Two-word prefixes: "电子BOM", "大配管BOM"
    if len(parts) >= 2 and parts[0] in ('电子', '大配管'):
        return ' '.join(parts[:2])
    return parts[0]


def group_diffs_by_parent(diff_rows):
    """Group diff rows by parent component — leaf-level, P3EM-centric view.

    Only P3EM parent nodes that directly contain leaf-diff items appear as
    group headers.  H5F parent nodes are hidden, but their leaf-level items
    are preserved and merged into the functionally-matching P3EM parent group.
    Intermediate assembly nodes (PNs that are themselves parent_pn of other
    diff rows) are excluded — they reflect structural hierarchy changes rather
    than material-level diffs.

    Flow:
      1. Collect all parent_pn values to identify leaf vs intermediate PNs.
      2. Partition all diff items: leaves go into groups by *their* direct
         parent_pn; intermediate PNs are dropped.
      3. Classify each leaf-group as P3EM (ADD-only), H5F (DEL-only), or MOD.
      4. Merge H5F leaf-groups into P3EM leaf-groups by functional-key matching.
      5. Assemble final result (P3EM + MOD + unmatched-H5F-safety-net).
    """
    # ── Step 1: Build the parent_pn look-up set and parent_name map ──
    parent_pn_set = set()
    parent_name_map = {}
    for dr in diff_rows:
        pp = (dr.get('parent_pn') or '').strip()
        pn = (dr.get('pn') or '').strip()
        if pp:
            parent_pn_set.add(pp)
            if pp not in parent_name_map:
                parent_name_map[pp] = dr.get('parent_name') or ''

    # ── Step 2: Partition leaves by their direct parent_pn ──
    # A PN is a "diff-leaf" when it never appears as parent_pn of any other
    # diff row.  Intermediate nodes (their PN ∈ parent_pn_set) are structural
    # hierarchy artifacts — they are dropped from the item list.
    leaf_groups = {}   # parent_pn → {'parent_pn': .., 'parent_name': .., adds/dels/mods}
    skipped_intermediate = []

    for dr in diff_rows:
        pn = (dr.get('pn') or '').strip()
        if pn in parent_pn_set:
            skipped_intermediate.append(pn)
            continue   # intermediate → skip

        # Leaf item — group by its direct parent
        pk = dr['parent_pn'] or '__UNKNOWN__'
        # Skip P1C* top-level parents
        if pk.upper().startswith('P1C'):
            continue

        if pk not in leaf_groups:
            leaf_groups[pk] = {
                'parent_pn': dr['parent_pn'],
                'parent_name': parent_name_map.get(pk, dr.get('parent_name') or ''),
                'adds': [], 'dels': [], 'mods': [],
            }
        if dr['type'] == 'added':
            leaf_groups[pk]['adds'].append(dr)
        elif dr['type'] == 'removed':
            leaf_groups[pk]['dels'].append(dr)
        else:
            leaf_groups[pk]['mods'].append(dr)

    # Filter out groups that ended up empty (shouldn't happen, but safety first)
    raw_groups = [g for g in leaf_groups.values()
                  if g['adds'] or g['dels'] or g['mods']]

    # ── Step 3: Classify each leaf-group as P3EM / H5F / MOD ──
    p3em_groups = []
    h5f_groups = []
    mod_groups = []
    for g in raw_groups:
        has_add = bool(g['adds'])
        has_del = bool(g['dels'])
        has_mod = bool(g['mods'])
        if has_add and not has_del and not has_mod:
            p3em_groups.append(g)
        elif has_del and not has_add and not has_mod:
            h5f_groups.append(g)
        else:
            mod_groups.append(g)

    # ── Step 4: Build P3EM functional-key index ──
    p3em_key_map = {}
    for pg in p3em_groups:
        key = _get_functional_key(pg['parent_name'])
        if key:
            p3em_key_map[key] = pg

    # ── Step 5: Merge H5F leaf-groups → P3EM by functional key ──
    unmatched_h5f = []
    for hg in h5f_groups:
        key = _get_functional_key(hg['parent_name'])
        matched = p3em_key_map.get(key)
        if matched:
            # All H5F leaf items go into the matched P3EM group
            for item in hg['dels']:
                if item['pn'] == hg['parent_pn']:
                    continue   # safety: skip H5F parent self-reference
                matched['dels'].append(item)
        else:
            unmatched_h5f.append(hg)

    # ── Step 6: Assemble final result ──
    # Sort groups by the earliest line_no from any item, preserving
    # the original BOM line_no order (not alphabetical by parent_pn).
    final_groups = p3em_groups + mod_groups + unmatched_h5f
    for g in final_groups:
        all_items = g['adds'] + g['dels'] + g['mods']
        g['_min_line_no'] = min(
            (it.get('line_no', 9999) for it in all_items),
            default=9999
        )
    final_groups.sort(key=lambda g: g['_min_line_no'])
    # Clean up internal key
    for g in final_groups:
        g.pop('_min_line_no', None)
    return final_groups


# ── Word Generation ────────────────────────────────────────


def _ensure_template():
    """如果模板文件不存在，自动创建一个基础模板。"""
    if os.path.exists(TEMPLATE_PATH):
        return
    doc = Document()
    # 段落0: 表单编号
    p0 = doc.add_paragraph()
    p0.add_run('SKY-RKZXS-07')
    # 段落1: 标题
    p1 = doc.add_paragraph()
    p1.add_run('更 改 通 知 单')
    # 表格: 6行 x 10列
    table = doc.add_table(rows=6, cols=10)
    # 第3行第1列作为更改内容区域
    # （模板填充逻辑依赖此结构）
    os.makedirs(os.path.dirname(TEMPLATE_PATH), exist_ok=True)
    doc.save(TEMPLATE_PATH)


def generate_change_notice(task_id: int, output_name: str = None, db_path: str = None,
                            order_no: str = None, stage: str = None, quantity: str = None,
                            drafter: str = None, reviewer: str = None):
    """Generate change notice by filling the official template.

    Parameters:
      order_no: 订单号
      stage:    阶段   (DVT/PVT/MP)
      quantity: 数量
      drafter:  拟制人
      reviewer: 审核人
    """
    _ensure_template()  # 确保模板存在
    if db_path is None:
        db_path = os.path.join(PROJECT_ROOT, 'data', 'bom_compare.db')
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row

    task = conn.execute('SELECT * FROM comparison_task WHERE id=?', (task_id,)).fetchone()
    if not task:
        raise ValueError(f'Task #{task_id} not found')

    src_row = conn.execute(
        'SELECT bom_name, bom_number FROM bom_header WHERE id=?',
        (task['source_bom_id'],)
    ).fetchone()
    tgt_row = conn.execute(
        'SELECT bom_name, bom_number FROM bom_header WHERE id=?',
        (task['target_bom_id'],)
    ).fetchone()

    if not src_row or not tgt_row:
        raise ValueError('BOM 数据不完整，请重新上传 BOM 文件')

    # 顶层BOM号：从 bom_item 查顶层物料的 part_number（最准确）
    src_top = _get_top_bom_number(conn, task['source_bom_id'])
    tgt_top = _get_top_bom_number(conn, task['target_bom_id'])
    if not src_top:
        src_top = src_row['bom_name'] or 'N/A'
    if not tgt_top:
        tgt_top = tgt_row['bom_name'] or 'N/A'

    machine_core = _extract_core(tgt_row['bom_number'] or tgt_top)
    # 从目标BOM的 bom_number 提取短机型名用于表头显示和文件名
    tgt_model = _extract_model(tgt_row['bom_number'] or tgt_top)

    diff_rows = get_diff_rows(conn, task_id,
                              source_bom_id=task['source_bom_id'],
                              target_bom_id=task['target_bom_id'])
    conn.close()

    today_str = datetime.now().strftime('%Y-%-m-%-d') if os.name != 'nt' else datetime.now().strftime('%Y/%#m/%#d')

    # Header fields — caller must supply valid values
    if not order_no:
        raise ValueError('订单号不能为空')
    if not stage:
        raise ValueError('阶段不能为空')
    if quantity is None:
        quantity = '1'

    # ── Load template ──
    doc = Document(TEMPLATE_PATH)

    # ── Update header paragraphs ──
    # P0: form number
    _clear_paragraph_runs(doc.paragraphs[0])
    run = doc.paragraphs[0].add_run('SKY-RKZXS-07')
    run.font.size = Pt(12)
    _set_font(run)

    # P1: title
    _clear_paragraph_runs(doc.paragraphs[1])
    run = doc.paragraphs[1].add_run('更 改 通 知 单')
    run.font.size = Pt(22)
    run.font.bold = True
    _set_font(run)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    doc.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Fill table header cells ──
    table = doc.tables[0]
    _fill_template_header(table, machine_core, today_str,
                          src_top, tgt_model,
                          src_label=src_row['bom_number'] or src_row['bom_name'],
                          tgt_label=tgt_row['bom_number'] or tgt_row['bom_name'],
                          order_no=order_no, stage=stage, quantity=quantity,
                          drafter=drafter, reviewer=reviewer)

    # ── Build content inside template table's "更改内容" cell ──
    # Table Row 3, Column 1 is the wide merged cell (span=9) marked "更改内容"
    table.rows[3].cells[1].merge(table.rows[3].cells[9])
    content_cell = table.rows[3].cells[1]
    groups = group_diffs_by_parent(diff_rows)
    _build_content_body(content_cell, groups)

    # ── Save ──
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if output_name is None:
        output_name = f'{order_no or "XX"}_{machine_core}_{tgt_model}_{quantity or "1"}台_{stage or "XX"}_整机更改通知单'
    output_path = os.path.join(OUTPUT_DIR, f'{output_name}.docx')
    doc.save(output_path)
    return output_path


def _clear_paragraph_runs(para):
    """Remove all runs from a paragraph."""
    for r in para.runs:
        r._element.getparent().remove(r._element)



def _set_cell_text(cell, text):
    """Set value cell text, matching template style (宋体 14pt centered)."""
    for p in cell.paragraphs:
        for r in p.runs:
            r._element.getparent().remove(r._element)

    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    _set_font(run)


def _fill_template_header(table, machine_core, today_str,
                         src_top, tgt_top,
                         src_label='', tgt_label='',
                         order_no='2606002KL', stage='DVT', quantity='1',
                         drafter=None, reviewer=None):
    """Fill template header rows (rows 0–5)."""

    # ── Quantity formatter: show as int if whole number ──
    def _fmt_qty(q):
        try:
            v = float(q)
            return str(int(v)) if v == int(v) else str(v)
        except (TypeError, ValueError):
            return str(q) if q else '1'

    # ── Row 0: 机芯 + 日期 ──
    _set_cell_text(table.rows[0].cells[3], machine_core)
    _set_cell_text(table.rows[0].cells[8], today_str)

    # ── Row 1: 机型 + 订单号 ──
    _set_cell_text(table.rows[1].cells[3], tgt_top)
    _set_cell_text(table.rows[1].cells[8], order_no)

    # ── Row 2: 阶段 + 数量 ──
    _set_cell_text(table.rows[2].cells[3], stage)
    _set_cell_text(table.rows[2].cells[8], quantity)

    # ── Row 4: 说明 ──
    note_text = (f'此份差异核对结果来源于 {src_label}（源BOM）与 {tgt_label}（目标BOM）'
                 f'的{stage}阶段比对，仅适用于 {tgt_label}（目标机型）使用。')
    note_cell = table.rows[4].cells[1]
    for p in list(note_cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = note_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(note_text)
    run.font.size = Pt(10)
    _set_font(run)

    # ── Row 5: 拟制 + 审核 ──
    if drafter:
        _set_cell_text(table.rows[5].cells[2], drafter)
    if reviewer:
        _set_cell_text(table.rows[5].cells[9], reviewer)


def _build_content_body(content_cell, groups):
    """Build component group content inside the template table's '更改内容' cell.

    All group headers and item lines are written as paragraphs within the
    designated merged cell (Table Row 3, Column 1).  This keeps all data
    strictly inside the template's existing table boundaries — no content
    spills outside as document-level paragraphs.
    """
    # ── Clear existing placeholder paragraphs in the cell ──
    # Remove all existing paragraphs to avoid blank lines before content
    for p in list(content_cell.paragraphs):
        p._element.getparent().remove(p._element)

    # ── Disable cantSplit on content row so Word CAN split it across pages ──
    from docx.oxml.ns import qn
    row3_tr = content_cell._tc.getparent()
    row3_trPr = row3_tr.find(qn('w:trPr'))
    if row3_trPr is not None:
        cant = row3_trPr.find(qn('w:cantSplit'))
        if cant is not None:
            row3_trPr.remove(cant)

    # ── Quantity formatter ──
    def _fmt_qty(q):
        try:
            v = float(q)
            return str(int(v)) if v == int(v) else str(v)
        except (TypeError, ValueError):
            return str(q) if q else '1'

    def _add_cell_para(text, font_size=Pt(10), bold=False, color='333333',
                       page_break_before=False):
        """Add a paragraph inside the content cell with consistent styling."""
        p = content_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.widow_control = False
        p.paragraph_format.keep_together = False
        if page_break_before:
            p.paragraph_format.page_break_before = True
        run = p.add_run(text)
        run.font.size = font_size
        run.font.bold = bold
        _set_font(run)
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        return p

    def _add_group_header(g):
        comp_pn = g.get('parent_pn', '')
        raw_name = g.get('short_name', g['parent_name'])
        # 用功能键截取核心组件名，去掉尾部技术参数（电压/功率/电流等）
        key = _get_functional_key(raw_name)
        comp_name = key if key and len(key) >= 2 else raw_name.split()[0] if raw_name.split() else raw_name
        _add_cell_para_counted(
            f'在{comp_pn} {comp_name}里',
            Pt(11), bold=True, color='1E40AF')

    def _add_item_line(prefix, pn, name, qty_text, ref=''):
        """Render one item line; if ref has >5 designators, wrap to continuation lines
        with left_indent calculated from first-line text width."""
        base = f'{pn} {name}     {qty_text}'
        prefix_str = f'{prefix}: ' if prefix else '     '
        first_line = prefix_str + base

        if not ref:
            _add_cell_para_counted(first_line, Pt(10), color='334155')
            return

        ref_parts = ref.split()
        ref_prefix = first_line + '  | '

        # Calculate indent width: ASCII ~5.5pt, CJK ~11pt in 10pt font
        ascii_cnt = sum(1 for c in ref_prefix if ord(c) <= 127)
        cjk_cnt = len(ref_prefix) - ascii_cnt
        indent_pt = ascii_cnt * 5.5 + cjk_cnt * 11.0

        # First line: up to 5 refs
        chunk0 = ref_parts[:5]
        _add_cell_para_counted(ref_prefix + ' '.join(chunk0), Pt(10), color='334155')

        # Continuation lines: 5 refs per line, indented to align with first ref
        for i in range(5, len(ref_parts), 5):
            chunk = ref_parts[i:i+5]
            p = _add_cell_para_counted(' '.join(chunk), Pt(10), color='334155')
            p.paragraph_format.left_indent = Pt(indent_pt)

    def _add_spacer():
        _add_cell_para_counted('', Pt(4))

    # ── Page-break logic: after every 28 lines, set page_break_before
    #     on the NEXT paragraph. Word respects paragraph-level page breaks
    #     even inside merged cells.
    LINES_PER_PAGE = 28
    line_counter = 0
    need_page_break = False

    def _add_cell_para_counted(text, font_size=Pt(10), bold=False, color='333333'):
        nonlocal line_counter, need_page_break
        line_counter += 1
        p = _add_cell_para(text, font_size=font_size, bold=bold, color=color,
                           page_break_before=need_page_break)
        return p
        need_page_break = (line_counter % LINES_PER_PAGE == 0)
        return p

    for g in groups:
        if not (g['adds'] or g['dels'] or g['mods']):
            continue

        comp_pn = g.get('parent_pn', '')
        if not comp_pn or comp_pn == '__UNKNOWN__':
            continue

        _add_group_header(g)

        # Collect all items into ADD and DEL buckets
        add_lines = []
        del_lines = []

        # Pure ADD items
        for item in g['adds']:
            qty = _fmt_qty(item['qty'])
            ref = item.get('ref', '')
            add_lines.append((item['pn'], item['name'], f'{qty}PC', ref))

        # Pure DEL items
        for item in g['dels']:
            qty = _fmt_qty(item['qty'])
            ref = item.get('ref', '')
            del_lines.append((item['pn'], item['name'], f'{qty}PC', ref))

        # MOD items — quantity changes → net ADD or DEL
        #              reference changes → split into ADD/DEL lines
        # 如果同一个 PN 同时有数量和位号变更，则位号 ADD/DEL 优先，
        # 数量行不再单独显示（位号不变的数量差无法拆分到具体位号）。
        mod_pns_with_ref = set()
        for m in g['mods']:
            if m.get('diff_category') == 'reference':
                mod_pns_with_ref.add(m['pn'])

        for m in g['mods']:
            is_qty = m.get('diff_category') == 'quantity'
            is_ref = m.get('diff_category') == 'reference'
            pn_key = m['pn']
            if is_qty:
                if pn_key in mod_pns_with_ref:
                    continue  # 位号变更已覆盖，不重复显示纯数量差
                try:
                    old_q = float(m.get('old_qty', 0) or 0)
                    new_q = float(m.get('new_qty', 0) or 0)
                except (TypeError, ValueError):
                    old_q = new_q = 0
                net = new_q - old_q
                if abs(net) < 0.001:
                    continue
                net_str = _fmt_qty(abs(net))
                if net > 0:
                    add_lines.append((m['pn'], m['name'], f'{net_str}PC', m.get('ref', '')))
                else:
                    del_lines.append((m['pn'], m['name'], f'{net_str}PC', m.get('ref', '')))
            elif is_ref:
                ref_added_str = (m.get('new_qty') or '').strip()
                ref_removed_str = (m.get('old_qty') or '').strip()
                if ref_added_str:
                    ref_added_list = [r for r in ref_added_str.split() if r]
                    cnt = len(ref_added_list)
                    add_lines.append((m['pn'], m['name'], f'{cnt}PC', ref_added_str))
                if ref_removed_str:
                    ref_removed_list = [r for r in ref_removed_str.split() if r]
                    cnt = len(ref_removed_list)
                    del_lines.append((m['pn'], m['name'], f'{cnt}PC', ref_removed_str))
            else:
                old_val = str(m.get('old_qty', '')).strip()
                new_val = str(m.get('new_qty', '')).strip()
                del_lines.append((m['pn'], m['name'], f'{old_val}\u2192{new_val}', m.get('ref', '')))

        # Render: ADD section first, then DEL section
        for i, (pn, nm, qt, ref) in enumerate(add_lines):
            _add_item_line('ADD' if i == 0 else '', pn, nm, qt, ref)
        for i, (pn, nm, qt, ref) in enumerate(del_lines):
            _add_item_line('DEL' if i == 0 else '', pn, nm, qt, ref)

        _add_spacer()


def _get_top_bom_number(conn, bom_id):
    """获取顶层BOM号（bom_item 中 level 最小的物料的 part_number）。"""
    row = conn.execute(
        'SELECT part_number FROM bom_item WHERE bom_id=? ORDER BY level ASC, line_no ASC LIMIT 1',
        (bom_id,)
    ).fetchone()
    return (row['part_number'] or '').strip() if row else ''


def _extract_model(bom_code):
    """Extract short model name from BOM code.

    Supports two formats:
      P1C{model}{core}{suffix}  → P1C85V68HP7T871001 → 85V68HP
      PCC{model}{core}{suffix}  → PCCF27G56M1MEC3000 → F27G56M
    """
    import re
    # PCC format (newer SAP exports): PCC + model + core + 3-digit suffix
    m = re.match(r'^PCC(.+?)(\d+[A-Z]+\d+)(\d{3})$', bom_code)
    if m:
        return m.group(1)
    # P1C format: P1C + model + core + 3-digit suffix
    m = re.match(r'^P1C(.+?)(\d+[A-Z]+\d+)(\d{3})$', bom_code)
    if m:
        return m.group(1)
    # Fallback: find first 3-digit + alphanumeric chunk
    m = re.search(r'(\d{3}[A-Z0-9]+)', bom_code)
    return m.group(1) if m else bom_code


def _extract_core(bom_code):
    """Extract machine core from full BOM code.

    Supports two formats:
      P1C{model}{core}{suffix}  → P1C100H5FP8R713002 → 8R713
      PCC{model}{core}{suffix}  → PCCF27G56M1MEC3000 → 1MEC3
    """
    import re
    # PCC format (newer SAP exports)
    m = re.match(r'^PCC(.+?)(\d+[A-Z]+\d+)(\d{3})$', bom_code)
    if m:
        return m.group(2)
    # P1C format
    m = re.match(r'^P1C(.+?)(\d+[A-Z]+\d+)(\d{3})$', bom_code)
    if m:
        return m.group(2)
    # Fallback
    m = re.search(r'(\d+[A-Z]+\d+)(\d{3})$', bom_code)
    if m:
        return m.group(1)
    return '8R713'


# ── Excel Export ────────────────────────────────────────────

def generate_change_notice_excel(task_id: int, output_name: str = None, db_path: str = None,
                                  drafter: str = None, reviewer: str = None):
    """Generate a professional Excel change notice matching the official form layout."""
    if db_path is None:
        db_path = os.path.join(PROJECT_ROOT, 'data', 'bom_compare.db')
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise ImportError('openpyxl required. pip install openpyxl')

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    task = conn.execute('SELECT * FROM comparison_task WHERE id=?', (task_id,)).fetchone()
    if not task:
        raise ValueError(f'Task #{task_id} not found')
    diff_rows = get_diff_rows(conn, task_id,
                              source_bom_id=task['source_bom_id'],
                              target_bom_id=task['target_bom_id'])
    conn.close()

    # 顶层BOM号：从 bom_item 查顶层物料 part_number
    conn2 = sqlite3.connect(db_path); conn2.row_factory = sqlite3.Row
    sn = conn2.execute('SELECT bom_name, bom_number FROM bom_header WHERE id=?', (task['source_bom_id'],)).fetchone()
    tn = conn2.execute('SELECT bom_name, bom_number FROM bom_header WHERE id=?', (task['target_bom_id'],)).fetchone()
    src_model = _get_top_bom_number(conn2, task['source_bom_id']) or (sn['bom_name'] if sn else 'SRC')
    tgt_model = _get_top_bom_number(conn2, task['target_bom_id']) or (tn['bom_name'] if tn else 'TGT')
    machine_core = _extract_core((tn['bom_number'] or '') if tn else '8R713')
    conn2.close()

    added_c = sum(1 for r in diff_rows if r['type'] == 'added')
    removed_c = sum(1 for r in diff_rows if r['type'] == 'removed')
    modified_c = sum(1 for r in diff_rows if r['type'] == 'modified')
    today_str = datetime.now().strftime('%Y-%-m-%-d') if os.name != 'nt' else datetime.now().strftime('%Y/%#m/%#d')

    # ── Workbook setup ──
    wb = Workbook()
    ws = wb.active
    ws.title = '整机清机更改通知单'

    # Colors
    blue_fill     = PatternFill(start_color='1E40AF', end_color='1E40AF', fill_type='solid')
    white_font    = Font(name='宋体', bold=True, color='FFFFFF', size=10)
    thin_border   = Border(
        left=Side('thin', '999999'), right=Side('thin', '999999'),
        top=Side('thin', '999999'), bottom=Side('thin', '999999'))
    no_border     = Border()
    center_a      = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_a        = Alignment(horizontal='left', vertical='center', wrap_text=True)
    right_a       = Alignment(horizontal='right', vertical='center')
    title_font    = Font(name='宋体', bold=True, size=16, color='1E40AF')
    subtitle_font = Font(name='宋体', bold=True, size=12)
    info_font     = Font(name='宋体', size=10)
    small_font    = Font(name='宋体', size=9)
    gray_small    = Font(name='宋体', size=8, color='999999')
    bold_font     = Font(name='宋体', bold=True, size=10)
    header_font   = Font(name='宋体', bold=True, color='FFFFFF', size=10)

    green_fill  = PatternFill(start_color='DCFCE7', end_color='DCFCE7', fill_type='solid')
    red_fill    = PatternFill(start_color='FEF2F2', end_color='FEF2F2', fill_type='solid')
    amber_fill  = PatternFill(start_color='FEF9C3', end_color='FEF9C3', fill_type='solid')
    green_font  = Font(name='宋体', color='16A34A', size=9)
    red_font    = Font(name='宋体', color='DC2626', size=9)
    amber_font  = Font(name='宋体', color='CA8A04', size=9)
    light_fill  = PatternFill(start_color='F8FAFC', end_color='F8FAFC', fill_type='solid')

    # ── Page setup ──
    ws.sheet_view.zoomScale = 90
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = 9
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.4
    ws.page_margins.bottom = 0.4

    # ── Row 1: Title ──
    ws.merge_cells('A1:H1')
    c = ws['A1']
    c.value = f'SKY-RKZXS-07          整 机 清 机 更 改 通 知 单'
    c.font = title_font
    c.alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[1].height = 36

    # ── Row 2: Info bar ──
    ws.merge_cells('A2:B2')
    ws['A2'].value = f'机芯：{machine_core}'
    ws['A2'].font = subtitle_font; ws['A2'].alignment = left_a
    ws.merge_cells('C2:D2')
    ws['C2'].value = f'机型：{src_model} → {tgt_model}'
    ws['C2'].font = subtitle_font; ws['C2'].alignment = left_a
    ws.merge_cells('E2:F2')
    ws['E2'].value = f'日期：{today_str}'
    ws['E2'].font = subtitle_font; ws['E2'].alignment = right_a
    ws.merge_cells('G2:H2')
    ws['G2'].value = f'阶段：整机清机'
    ws['G2'].font = subtitle_font; ws['G2'].alignment = right_a
    ws.row_dimensions[2].height = 24

    # ── Row 3: Summary ──
    ws.merge_cells('A3:H3')
    ws['A3'].value = (f'比对汇总：新增 {added_c} 项 | 删除 {removed_c} 项 | 变更 {modified_c} 项 | '
                      f'合计 {len(diff_rows)} 项差异    |    源机型：{src_model}    |    目标机型：{tgt_model}')
    ws['A3'].font = Font(name='宋体', size=9, italic=True, color='64748B')
    ws['A3'].alignment = left_a
    ws['A3'].fill = light_fill
    ws.row_dimensions[3].height = 22

    # ── Row 4: Blank spacer ──
    ws.row_dimensions[4].height = 6

    # ── Row 5: Column headers ──
    diff_headers = ['序号', '差异类型', '物料编码', '物料名称', '规格', '原用量', '新用量', '位号']
    col_widths_excel = [5, 8, 20, 36, 18, 7, 7, 14]

    for ci, (h, w) in enumerate(zip(diff_headers, col_widths_excel), 1):
        cell = ws.cell(row=5, column=ci, value=h)
        cell.font = header_font
        cell.fill = blue_fill
        cell.alignment = center_a
        cell.border = thin_border
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.row_dimensions[5].height = 24

    # ── Data rows (starting Row 6) ──
    for di, dr in enumerate(diff_rows):
        row_num = di + 6
        dt = dr['type']

        if dt == 'added':
            fill = green_fill; font = green_font
            old_qty = '-'; new_qty = str(dr.get('qty', '1'))
        elif dt == 'removed':
            fill = red_fill; font = red_font
            old_qty = str(dr.get('qty', '1')); new_qty = '-'
        else:
            fill = amber_fill; font = amber_font
            old_qty = str(dr.get('old_qty', '')); new_qty = str(dr.get('new_qty', ''))

        values_excel = [di + 1, dr['type_label'], dr['pn'], dr['name'][:120],
                        '', old_qty, new_qty, '']

        for ci, val in enumerate(values_excel, 1):
            cell = ws.cell(row=row_num, column=ci, value=val)
            cell.font = font
            cell.fill = fill
            cell.border = thin_border
            cell.alignment = center_a if ci in (1, 2, 6, 7) else left_a
        ws.row_dimensions[row_num].height = 18

    # ── Signature section ──
    sig_row = len(diff_rows) + 7
    ws.merge_cells(f'A{sig_row}:H{sig_row}')
    ws.row_dimensions[sig_row].height = 6

    sig_row += 1
    ws.merge_cells(f'A{sig_row}:C{sig_row}')
    ws[f'A{sig_row}'].value = f'拟制：{drafter or "杨芮"}'
    ws[f'A{sig_row}'].font = bold_font
    ws[f'A{sig_row}'].alignment = left_a

    ws.merge_cells(f'E{sig_row}:H{sig_row}')
    ws[f'E{sig_row}'].value = f'审核：{reviewer or "程涛"}'
    ws[f'E{sig_row}'].font = bold_font
    ws[f'E{sig_row}'].alignment = right_a

    sig_row += 1
    ws.merge_cells(f'A{sig_row}:H{sig_row}')
    ws[f'A{sig_row}'].value = f'本通知单由 BOM 比对工具自动生成 | 生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    ws[f'A{sig_row}'].font = gray_small
    ws[f'A{sig_row}'].alignment = right_a

    # Freeze panes below header
    ws.freeze_panes = 'A6'

    # Auto-filter
    last_data_row = len(diff_rows) + 5
    ws.auto_filter.ref = f'A5:H{last_data_row}'

    # ── Save ──
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if output_name is None:
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_name = f'整机清机更改通知单_{src_model}_vs_{tgt_model}_{ts}'
    output_path = os.path.join(OUTPUT_DIR, f'{output_name}.xlsx')
    wb.save(output_path)
    return output_path


# ── CLI ─────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys
    tid = int(sys.argv[1]) if len(sys.argv) > 1 else 45
    fmt = sys.argv[2] if len(sys.argv) > 2 else 'docx'

    print(f'生成整机清机更改通知单...')
    print(f'  任务 ID: {tid}')
    if fmt == 'xlsx':
        path = generate_change_notice_excel(tid)
    else:
        path = generate_change_notice(tid)
    print(f'  输出文件: {path}')
    print(f'  完成!')
