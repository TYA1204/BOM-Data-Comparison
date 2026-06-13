"""Compare L1 component presence across documents."""
from docx import Document
import re, os

files = [
    ('094542 (09:45)', 'C:/Users/1005691/Downloads/整机清机更改通知单_100H5F_vs_100P3EM_20260612_094542.docx'),
    ('100207 (10:02)', 'C:/Users/1005691/Downloads/整机清机更改通知单_100H5F_vs_100P3EM_20260612_100207.docx'),
    ('102559 (10:25)', 'C:/Users/1005691/Downloads/整机清机更改通知单_100H5F_vs_100P3EM_20260612_102559.docx'),
    ('103428 (10:34) WRONG', 'C:/Users/1005691/Downloads/整机清机更改通知单_100H5F_vs_100P3EM_20260612_103428.docx'),
    ('153811 (15:38) CORRECT', 'reports/整机清机更改通知单_100H5F_vs_100P3EM_20260612_153811.docx'),
    ('155052 (15:50) OPT', 'reports/整机清机更改通知单_100H5F_vs_100P3EM_20260612_155052.docx'),
]

for label, fp in files:
    doc = Document(fp)
    
    # Extract document-level groups
    groups = []
    current_group = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        m_group = re.match(r'^在(.+?)··(.+?)里$', t)
        if m_group:
            current_group = {'pn': m_group.group(1), 'name': m_group.group(2), 'items': []}
            groups.append(current_group)
            continue
        m_pre = re.match(r'^(ADD|DEL|MOD):([\w\-]+)··(.+?)··(\S+)$', t)
        m_cont = re.match(r'^\s{5}([\w\-]+)··(.+?)··(\S+)$', t)
        if m_pre and current_group:
            current_group['items'].append({'prefix': m_pre.group(1), 'pn': m_pre.group(2)})
        elif m_cont and current_group:
            current_group['items'].append({'prefix': '', 'pn': m_cont.group(1)})
    
    # Extract table-cell content (old format)
    tb_items = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                ct = cell.text.strip()
                found = re.findall(r'(ADD|DEL|MOD):([\w\-]+)··', ct)
                tb_items.extend(found)
    
    # Check for P1C in any content
    has_p1c = any(g['pn'].startswith('P1C') for g in groups)
    has_p1c_table = any(t[1].startswith('P1C') for t in tb_items)
    all_pns = set()
    for g in groups:
        for item in g['items']:
            all_pns.add(item['pn'])
    
    total_items = sum(len(g['items']) for g in groups)
    l1_groups = [g for g in groups if g['pn'].startswith('P1C')]
    
    print(f'\n=== {label} ===')
    print(f'  Format: {len(groups)} doc-level groups, {total_items} items | {len(tb_items)} table items')
    print(f'  P1C in doc-level: {has_p1c} | P1C in table: {has_p1c_table}')
    if l1_groups:
        for g in l1_groups:
            pns = [i['pn'] for i in g['items']]
            print(f'  L1 group {g["pn"]}: {pns}')
    
    # Also show all group PNs
    group_pns = [g['pn'] for g in groups]
    if group_pns:
        print(f'  All group PNs: {group_pns}')
    
    # For table-cell format, show some content
    if tb_items:
        print(f'  Table cell PNs (sample): {tb_items[:10]}')
    
    # Check paragraphs count
    content_paras = [p.text.strip()[:100] for p in doc.paragraphs if p.text.strip()]
    pb_count = sum(1 for p in doc.paragraphs for r in p.runs if 'w:br' in r._element.xml and 'type="page"' in r._element.xml)
    print(f'  Content paragraphs: {len(content_paras)}, Page breaks: {pb_count}')
