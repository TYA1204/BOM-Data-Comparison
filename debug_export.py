"""Debug: generate Word and verify content is inside table cell."""
import sys
sys.path.insert(0, 'D:/BOM Data Comparison')

from app.services.change_notice import generate_change_notice

output_path = generate_change_notice(task_id=20)
print(f'Generated: {output_path}')

# Verify structure
from docx import Document
doc = Document(output_path)

print(f'\n表格数: {len(doc.tables)}')
print(f'段落数: {len(doc.paragraphs)}')

# Check: paragraphs should NOT contain diff content (only title stuff)
has_outside_content = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and '在N0' in text:
        has_outside_content = True
        print(f'  ⚠ OUTSIDE TABLE P[{i}]: {text[:60]}')
    elif text and ('ADD:' in text or 'DEL:' in text or 'MOD:' in text):
        has_outside_content = True
        print(f'  ⚠ OUTSIDE TABLE P[{i}]: {text[:60]}')

if not has_outside_content:
    print('OK: 无变更内容溢出到表格外部')

# Check: content cell should have the diff data
table = doc.tables[0]
content_cell = table.rows[3].cells[1]
cell_paragraphs = content_cell.paragraphs
print(f'\n"更改内容"单元格内段落数: {len(cell_paragraphs)}')
for i, p in enumerate(cell_paragraphs):
    text = p.text.strip()
    if text:
        print(f'  行{i}: {text[:80]}')

# Verify: no other document-level paragraphs beyond the expected ones
print(f'\n文档级段落（非表格内）:')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'  P[{i}]: {text[:80]}')
